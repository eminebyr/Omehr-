from __future__ import annotations

"""CI regresyon bariyerleri.

Bu araç, geliştirme sürecinde AYNI hata sınıfının BAĞIMSIZ olarak
birden çok kez (bazen 5-6 farklı dosyada) yeniden ortaya çıktığı 3
somut soruna karşı OTOMATİK, kod-inceleme-öncesi koruma sağlar:

1. Sabit kodlanmış e-posta adresleri — servis/web katmanında HİÇBİR
   e-posta adresi doğrudan string sabiti olarak YAZILAMAZ; her zaman
   çalışma zamanında veriden (Mail_Listesi, tenant kaydı) gelmelidir.
   Aksi halde çok kiracılı bir SaaS'ta bir kiracının bildirimleri
   başka bir kiracıya/kişiye gider (gerçekte 6+ kez bulundu).

2. Ana input dosyasının önbelleksiz doğrudan okunması —
   `pd.read_excel(INPUT, ...)` gibi çağrılar, services/
   cached_excel_reader.py'nin sunduğu paylaşımlı önbelleği atlar ve
   ölçülmüş, gerçek bir performans regresyonuna yol açar (gerçekte
   10+ dosyada bulundu, ~0.5 sn/çağrı maliyetliydi).

3. Sürüm numarası senkron kayması — services/version.py'deki
   APP_VERSION, belgelenen "canlı" referans dosyalarının HEPSİNDE
   aynı olmalıdır; aksi halde kullanıcıya yanlış sürüm bilgisi
   gösterilir (gerçekte 4-5 kez yaşandı).
"""

import ast
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

# ------------------------------------------------------------------
# 1) Sabit kodlanmış e-posta adresleri
# ------------------------------------------------------------------
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")

# Bilinçli, meşru istisnalar: örnek/şablon amaçlı, gerçek bir alıcıya
# ASLA gitmeyen yer tutucular. Bu liste KISA tutulmalı — her ekleme
# gerekçelendirilmeli.
_EMAIL_ALLOWLIST = {
    "kullanici@ornek.com", "ornek@ornek.com", "test@test.com",
    "admin@example.com", "user@example.com",
}
_EMAIL_ALLOWED_DIRS = {"tests", "archive", "ORNEK_TEST_VERISI"}
_EMAIL_ALLOWED_FILES = {".env.example"}


def _sabit_email_ihlalleri(root: Path) -> list[str]:
    problems: list[str] = []
    for base in ("services", "web"):
        for path in sorted((root / base).rglob("*.py")):
            rel = path.relative_to(root)
            if any(part in _EMAIL_ALLOWED_DIRS for part in rel.parts):
                continue
            if path.name in _EMAIL_ALLOWED_FILES:
                continue
            try:
                tree = ast.parse(path.read_text(encoding="utf-8"), filename=str(path))
            except SyntaxError:
                continue
            for node in ast.walk(tree):
                if isinstance(node, ast.Constant) and isinstance(node.value, str):
                    for match in _EMAIL_RE.findall(node.value):
                        if match.casefold() in _EMAIL_ALLOWLIST:
                            continue
                        problems.append(
                            f"{rel}:{node.lineno}: sabit kodlanmış e-posta adresi bulundu: '{match}' — "
                            "e-posta adresleri HER ZAMAN çalışma zamanında veriden (Mail_Listesi/tenant "
                            "kaydı) okunmalı, kaynak koda gömülmemelidir."
                        )
    return problems


# ------------------------------------------------------------------
# 2) Önbelleksiz doğrudan Excel okuması
# ------------------------------------------------------------------
_READ_ALLOWED_FILES = {
    "cached_excel_reader.py",   # önbelleğin kendisi
    "master_data_admin.py",     # YAZMA öncesi taze okuma kasıtlı
    "excel_read_shim.py",       # pandas.read_excel'in kendisini monkeypatch eder
    "security.py",              # migrate_legacy_input: tek seferlik kurulum okuması
    "formula_bagimsiz_hesapla.py",  # ana motorun tek-seferlik parti hesaplaması
    "model_governance.py",      # düşük sıklıklı denetim raporu
    "input_excel_migration.py", # DB'ye BİR KEZ taşıma işlemi
    "report_pipeline.py",       # main.py'nin JUST-ÜRETİLMİŞ çıktı raporunu tek
                                 # seferlik doğrulaması — ana input dosyası DEĞİL,
                                 # her çalıştırmada zaten taze; önbellek değer katmaz
}
_READ_TARGET_NAMES = {"INPUT", "input_path", "hedef", "INPUT_PATH", "path"}


def _onbelleksiz_okuma_ihlalleri(root: Path) -> list[str]:
    problems: list[str] = []
    for base in ("services", "web"):
        for path in sorted((root / base).rglob("*.py")):
            if path.name in _READ_ALLOWED_FILES:
                continue
            rel = path.relative_to(root)
            if any(part in {"tests", "archive"} for part in rel.parts):
                continue
            try:
                tree = ast.parse(path.read_text(encoding="utf-8"), filename=str(path))
            except SyntaxError:
                continue
            for node in ast.walk(tree):
                if not (isinstance(node, ast.Call) and isinstance(node.func, ast.Attribute)):
                    continue
                if node.func.attr != "read_excel":
                    continue
                if not node.args:
                    continue
                first = node.args[0]
                isim = first.id if isinstance(first, ast.Name) else None
                if isim in _READ_TARGET_NAMES:
                    problems.append(
                        f"{rel}:{node.lineno}: '{isim}' üzerinde önbelleksiz pd.read_excel(...) çağrısı — "
                        "services.cached_excel_reader.read_sheet_cached(...) kullanın (ölçülmüş performans "
                        "regresyonu: ~0.5 sn/çağrı). Bu dosya gerçekten önbelleksiz taze okuma gerektiriyorsa "
                        "tools/check_regression_guards.py'deki _READ_ALLOWED_FILES listesine GEREKÇEYLE ekleyin."
                    )
    return problems


# ------------------------------------------------------------------
# 3) Sürüm numarası senkron kayması
# ------------------------------------------------------------------
_VERSION_FILES = ["00_OKU_CURRENT.txt", "SURUM_NOTLARI.md", "DOGRULAMA_RAPORU.md", "KISA_KURULUM.txt"]


def _surum_uyumsuzluklari(root: Path) -> list[str]:
    problems: list[str] = []
    version_py = root / "services" / "version.py"
    if not version_py.exists():
        return problems
    match = re.search(r'APP_VERSION\s*=\s*"([^"]+)"', version_py.read_text(encoding="utf-8"))
    if not match:
        return [f"services/version.py: APP_VERSION bulunamadı — sürüm kontrolü yapılamıyor."]
    surum = match.group(1)

    for dosya_adi in _VERSION_FILES:
        dosya = root / dosya_adi
        if not dosya.exists():
            continue
        icerik = dosya.read_text(encoding="utf-8", errors="ignore")
        if surum not in icerik:
            problems.append(
                f"{dosya_adi}: services/version.py'deki APP_VERSION ('{surum}') bu dosyada geçmiyor — "
                "sürüm senkron kaymış olabilir."
            )

    kilavuz = root / "KULLANICI_KILAVUZU.docx"
    if kilavuz.exists():
        try:
            import docx
            d = docx.Document(str(kilavuz))
            metin = "\n".join(p.text for p in d.paragraphs)
            for section in d.sections:
                metin += "\n" + "\n".join(p.text for p in section.footer.paragraphs)
            if surum not in metin:
                problems.append(
                    f"KULLANICI_KILAVUZU.docx: services/version.py'deki APP_VERSION ('{surum}') "
                    "kılavuzun gövde/alt bilgi metninde geçmiyor — sürüm senkron kaymış olabilir."
                )
        except Exception as exc:
            problems.append(f"KULLANICI_KILAVUZU.docx: sürüm kontrolü yapılamadı ({exc}).")

    return problems


def violations(root: Path = ROOT) -> list[str]:
    return (
        _sabit_email_ihlalleri(root)
        + _onbelleksiz_okuma_ihlalleri(root)
        + _surum_uyumsuzluklari(root)
    )


def main() -> int:
    problems = violations()
    if problems:
        print("REGRESYON BARİYERİ İHLALLERİ")
        for p in problems:
            print(" -", p)
        return 1
    print("Regression guards: OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
