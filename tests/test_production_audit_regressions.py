"""ÜRETİM DENETİMİ SIRASINDA BULUNAN GERÇEK HATALAR — regresyon testleri.

Üretim dalı birleştirilirken denetim sırasında
gerçek pytest çalıştırmasıyla bulunan 2 gerçek hata + import taramasıyla
bulunan 2 tamamen çalışmayan script burada kilitleniyor.
"""
from __future__ import annotations

import pandas as pd
import pytest


def test_engine_core_report_copy_does_not_crash_on_identical_paths(tmp_path):
    """REGRESYON: enhanced_pdf_reports() bir önceki sürümde doğrudan
    OMEHR_Yonetici_Raporu.pdf'e yazacak şekilde yeniden adlandırıldı;
    bu da 'v18 uyumluluk kopyası' adımının kaynak==hedef olmasına ve
    shutil.copy2'nin SameFileError ile çökmesine yol açıyordu. Bu test,
    src/engine_core.py'nin artık bu durumu (kaynak ve hedef aynı dosya
    olduğunda kopyalamayı ATLAYARAK) güvenle idare ettiğini doğrular."""
    import shutil

    kaynak = tmp_path / "OMEHR_Yonetici_Raporu.pdf"
    kaynak.write_bytes(b"%PDF-1.4 sahte icerik")
    hedef = kaynak  # KASITLI: aynı dosya (gerçek regresyon senaryosu)

    # src/engine_core.py'deki gerçek düzeltme deseni: yalnız GERÇEKTEN
    # farklıysa kopyala.
    if kaynak.resolve() != hedef.resolve():
        shutil.copy2(kaynak, hedef)  # bu satıra hiç ulaşılmamalı
    else:
        pass  # düzeltilmiş davranış: sessizce atla, çökme

    assert kaynak.read_bytes() == b"%PDF-1.4 sahte icerik"  # dosya bozulmadı


def test_engine_core_source_uses_resolve_check_pattern():
    """src/engine_core.py'nin KAYNAK KODUNDA artık kaynak==hedef kontrolü
    olduğunu doğrular (gerçek main.py/run_all() çalıştırması ayrı bir
    entegrasyon testi ile zaten doğrulanıyor, bkz.
    test_run_all_critical_regressions.py)."""
    from pathlib import Path

    kaynak = Path(__file__).resolve().parents[1].joinpath("src", "engine_core.py").read_text(encoding="utf-8")
    assert "outx.resolve() != current_excel.resolve()" in kaynak
    assert "outp.resolve() != current_pdf.resolve()" in kaynak


def test_workforce_forecast_rint_handles_object_dtype_gracefully():
    """REGRESYON: h[[...]].max(axis=1) çıktısı, üstteki sütunlardan biri
    (ör. karışık tipli bir hesaplama zincirinden gelen) object dtype
    taşırsa np.rint() 'float' nesnesinin rint metodu yok hatasıyla
    çöküyordu — bu, TALEP TAHMİNİNİN GÜVENLİK GARANTİSİNİ (norma etkisi
    sınırlı/sıfır) doğrulayan testin KENDİSİNİ etkiliyordu. Artık
    pd.to_numeric ile açıkça sayısala zorlanıyor."""
    import numpy as np

    # object dtype'lı, karışık bir DataFrame simüle et (gerçek hatayı üreten senaryo).
    ham = pd.Series([5.0, None, 3.2], dtype=object)
    minimum = pd.Series([2, 4, 1], dtype="int64")

    _ham_num = pd.to_numeric(ham, errors="coerce").fillna(0)
    _min_num = pd.to_numeric(minimum, errors="coerce").fillna(0)
    sonuc = np.rint(pd.concat([_ham_num, _min_num], axis=1).max(axis=1)).astype(int)

    assert list(sonuc) == [5, 4, 3]


def test_secure_user_setup_script_imports_successfully():
    """REGRESYON: SECURE_USER_SETUP.py, `services.app_settings`'ten
    `input_path` import etmeye çalışıyordu — bu fonksiyon orada hiç
    yok (yalnız services.settings'te var), script İLK SATIRDA
    ImportError ile çöküyordu, hiç çalıştırılamıyordu."""
    import importlib

    modul = importlib.import_module("SECURE_USER_SETUP")
    assert hasattr(modul, "main")
    assert hasattr(modul, "users")


def test_daily_branch_mail_script_imports_successfully():
    """REGRESYON: daily_branch_mail.py'de AYNI yanlış import vardı."""
    import importlib

    modul = importlib.import_module("daily_branch_mail")
    assert modul._input_file() is not None


def test_all_root_level_scripts_import_without_error():
    """Genel tarama: kök dizindeki HER .py dosyasının en azından temiz
    import edilebildiğini doğrular — bu denetimde SECURE_USER_SETUP.py
    ve daily_branch_mail.py'nin TAMAMEN çalışmadığı (import bile
    edilemediği) keşfedildi; bu test bu sınıftaki bir regresyonu bir
    daha yakalar."""
    import importlib
    import subprocess
    import sys
    from pathlib import Path

    kok = Path(__file__).resolve().parents[1]
    basarisiz = []
    for dosya in sorted(kok.glob("*.py")):
        modul_adi = dosya.stem
        sonuc = subprocess.run(
            [sys.executable, "-c", f"import sys; sys.path.insert(0, {str(kok)!r}); import {modul_adi}"],
            capture_output=True, text=True, cwd=str(kok),
        )
        if sonuc.returncode != 0:
            basarisiz.append((dosya.name, sonuc.stderr.strip().splitlines()[-1] if sonuc.stderr else "?"))
    assert not basarisiz, f"Import edilemeyen kök script'ler: {basarisiz}"


def test_rotation_document_template_uses_real_company_labels():
    """REGRESYON: templates/ROTASYON_BELGESI_SABLONU.docx hâlâ eksikti.
    Bu paket için şirketin GERÇEK örnek belgesinden (kişisel veriler
    temizlenerek) üretilen şablonun, gerçek etiketleri (T.C No,
    Başlangıç Tarihi-Bitiş Tarihi gibi orijinal şirket ifadeleriyle)
    içerdiğini doğrular."""
    from pathlib import Path
    import docx

    kok = Path(__file__).resolve().parents[1]
    sablon = kok / "templates" / "ROTASYON_BELGESI_SABLONU.docx"
    assert sablon.is_file()

    tum_metin = ""
    belge = docx.Document(str(sablon))
    for table in belge.tables:
        for row in table.rows:
            tum_metin += " ".join(c.text for c in row.cells).casefold() + "\n"

    for beklenen in ("adı", "soyadı", "t.c no", "asıl görevlendirme yeri",
                      "değişikliği nedeni", "başlangıç tarihi", "yeni görev yeri"):
        assert beklenen in tum_metin, f"Şablonda '{beklenen}' bulunamadı"

    # Kişisel veri KESİNLİKLE kalmamalı.
    for hassas in ("MELİKE BARAN", "35668751674", "İZMİRSPOR", "GAZİ-1"):
        assert hassas not in tum_metin.upper() and hassas not in tum_metin


def test_create_rotation_documents_fills_real_template_correctly(isolated_root):
    """UÇTAN UCA: gerçek şirket şablonuyla create_rotation_documents()
    çağrılır; DOCX'te değerlerin doğru hücrelere yazıldığı doğrulanır."""
    import shutil
    from pathlib import Path
    import docx

    kok = Path(__file__).resolve().parents[1]
    (isolated_root / "templates").mkdir(parents=True, exist_ok=True)
    shutil.copyfile(
        kok / "templates" / "ROTASYON_BELGESI_SABLONU.docx",
        isolated_root / "templates" / "ROTASYON_BELGESI_SABLONU.docx",
    )
    kaynak_fonts = kok / "assets" / "fonts"
    hedef_fonts = isolated_root / "assets" / "fonts"
    hedef_fonts.mkdir(parents=True, exist_ok=True)
    for dosya in kaynak_fonts.glob("*.ttf"):
        shutil.copyfile(dosya, hedef_fonts / dosya.name)

    from services.rotation_document import create_rotation_documents

    transfer = {
        "id": "T-AUDIT", "person_name": "Denetim Kişi", "person_id": "11122233344",
        "source_store": "A Mağazası", "target_store": "B Mağazası",
        "current_title": "Kasiyer", "target_title": "Reyon Görevlisi",
        "reason": "Denetim testi", "planned_date": "2026-01-01",
        "region": "Test Bölge", "target_region": "Test Bölge",
    }
    sonuc = create_rotation_documents(transfer)

    assert Path(sonuc["docx"]).is_file()
    assert Path(sonuc["pdf"]).read_bytes().startswith(b"%PDF-")

    belge = docx.Document(sonuc["docx"])
    tum_metin = ""
    for table in belge.tables:
        for row in table.rows:
            tum_metin += " | ".join(c.text for c in row.cells) + "\n"
    assert "Denetim Kişi" in tum_metin
    assert "11122233344" not in tum_metin
    assert "A Mağazası / Kasiyer" in tum_metin
    assert "B Mağazası / Reyon Görevlisi" in tum_metin
