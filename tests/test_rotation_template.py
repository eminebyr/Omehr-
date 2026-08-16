from __future__ import annotations

from pathlib import Path
from docx import Document


def test_rotation_template_is_anonymized_and_fillable(isolated_root):
    import shutil
    source=Path(__file__).resolve().parents[1]/"templates"/"ROTASYON_BELGESI_SABLONU.docx"
    target=isolated_root/"templates"/source.name
    target.parent.mkdir(parents=True,exist_ok=True)
    shutil.copy2(source,target)
    font_source=Path(__file__).resolve().parents[1]/"assets"/"fonts"
    font_target=isolated_root/"assets"/"fonts"
    font_target.mkdir(parents=True,exist_ok=True)
    for name in ("DejaVuSans.ttf","DejaVuSans-Bold.ttf"):
        shutil.copy2(font_source/name,font_target/name)

    from services.rotation_document import create_rotation_documents
    result=create_rotation_documents({
        "id":"T-1","person_name":"DENEME PERSONEL","person_id":"P0001",
        "source_store":"AKEVLER","target_store":"GAZİ-1",
        "current_title":"REYON GÖREVLİSİ","target_title":"REYON GÖREVLİSİ",
        "reason":"NORM DENGELEME","planned_date":"04.08.2026",
    })
    doc=Document(result["docx"])
    text="\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "DENEME PERSONEL" in text
    assert "P0001" not in text
    assert "AKEVLER / REYON GÖREVLİSİ" in text
    assert "GAZİ-1 / REYON GÖREVLİSİ" in text
    assert Path(result["pdf"]).read_bytes().startswith(b"%PDF-")


def test_rotation_template_contains_no_original_personal_data():
    source=Path(__file__).resolve().parents[1]/"templates"/"ROTASYON_BELGESI_SABLONU.docx"
    doc=Document(source)
    text="\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "MELİKE BARAN" not in text
    assert "35668751674" not in text
