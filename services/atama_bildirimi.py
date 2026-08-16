from __future__ import annotations

"""ATAMA / GÖREV DEĞİŞİKLİĞİ BİLDİRİMİ oluşturur.

İşe giriş/çıkış bildirimlerinden FARKLI bir olay: aynı kişi aktif kalır,
yalnız Unvan (ve/veya Mağaza) değişir — bir terfi, rotasyon veya görev
değişikliği. `services/personnel_exit.py::update_personnel()` ile Unvan
değiştirildiğinde bu modül çağrılarak hem imzalanabilir bir DOCX/PDF
belge üretilir hem de ilgili taraflara e-posta gönderilir.

templates/ATAMA_BILDIRIMI_SABLONU.docx şablonu KASITLI olarak
kiracıdan bağımsızdır (hiçbir şirket adı gömülü değil) — {{SIRKET_ADI}}
yerine tenant_registry'deki firma adı veya genel bir ifade konur.
"""

import os
import re
import tempfile
import unicodedata
from datetime import datetime
from pathlib import Path

from docx import Document
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from services.pdf_compat import make_outlook_safe_pdf
from services.runtime_paths import runtime_root
from src.pdf_fonts import font as _pdf_font

def _template():
    from services.runtime_paths import runtime_root
    return runtime_root() / "templates" / "ATAMA_BILDIRIMI_SABLONU.docx"
def _output_dir():
    from services.runtime_paths import runtime_root
    return runtime_root() / "output" / "Atama_Bildirimleri"


def _safe_name(value: object) -> str:
    table = str.maketrans({"ç":"c","Ç":"C","ğ":"g","Ğ":"G","ı":"i","İ":"I","ö":"o","Ö":"O","ş":"s","Ş":"S","ü":"u","Ü":"U"})
    text = unicodedata.normalize("NFKD", str(value or "").translate(table))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return re.sub(r"[^A-Za-z0-9._-]+", "_", text).strip("_") or "atama"


def _tenant_display_name() -> str:
    """Kiracının GERÇEK ticari adını döner — sabit kodlanmış bir şirket
    adı KULLANILMAZ (çok kiracılı SaaS'ta yanlış olurdu)."""
    try:
        from services.tenant_context import current_tenant_id
        from services.tenant_registry import get_tenant
        bilgi = get_tenant(current_tenant_id())
        if bilgi and bilgi.get("name"):
            return str(bilgi["name"])
    except Exception:
        pass
    return "Şirketimiz"


def _replace_placeholders(paragraphs, mapping: dict[str, str]) -> None:
    for p in paragraphs:
        for run in p.runs:
            for key, val in mapping.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)
        # Word bazen tek bir metni birden fazla run'a böler; birleşik
        # paragraf metninde hâlâ yer tutucu kalırsa ilk run'a yaz.
        birlesik = "".join(r.text for r in p.runs)
        for key, val in mapping.items():
            if key in birlesik and p.runs:
                birlesik = birlesik.replace(key, val)
                for i, r in enumerate(p.runs):
                    r.text = birlesik if i == 0 else ""


def create_assignment_notice(atama: dict) -> dict[str, str]:
    """atama: isim_soyisim, yeni_pozisyon, yeni_magaza, onceki_pozisyon,
    onceki_magaza, tarih, onaylayan içerir."""
    _output_dir().mkdir(parents=True, exist_ok=True)
    isim = str(atama.get("isim_soyisim") or "")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = f"ATAMA_{_safe_name(isim)}_{stamp}"
    final_docx = _output_dir() / f"{stem}.docx"
    final_pdf = _output_dir() / f"{stem}.pdf"

    mapping = {
        "{{SIRKET_ADI}}": _tenant_display_name(),
        "{{AD_SOYAD}}": isim,
        "{{MAGAZA}}": str(atama.get("yeni_magaza") or ""),
        "{{TARIH}}": str(atama.get("tarih") or ""),
        "{{YENI_POZISYON}}": str(atama.get("yeni_pozisyon") or ""),
        "{{ONCEKI_POZISYON}}": str(atama.get("onceki_pozisyon") or ""),
        "{{ONCEKI_MAGAZA}}": str(atama.get("onceki_magaza") or ""),
        "{{ONAYLAYAN}}": str(atama.get("onaylayan") or ""),
    }

    with tempfile.TemporaryDirectory(prefix="atama_", dir=str(_output_dir())) as tmpdir:
        tmp = Path(tmpdir)
        tmp_docx = tmp / final_docx.name
        tmp_pdf = tmp / final_pdf.name

        if not _template().exists():
            raise FileNotFoundError(f"Atama şablonu bulunamadı: {_template()}")
        doc = Document(str(_template()))
        _replace_placeholders(doc.paragraphs, mapping)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _replace_placeholders(cell.paragraphs, mapping)
        doc.save(str(tmp_docx))
        if not tmp_docx.is_file() or tmp_docx.stat().st_size < 1000:
            raise IOError("Atama DOCX dosyası oluşturulamadı.")

        font = _pdf_font(bold=False)
        font_bold = _pdf_font(bold=True)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("atama_title", parent=styles["Heading1"], fontName=font_bold, fontSize=13, leading=17, alignment=TA_CENTER, spaceAfter=10, textColor=colors.HexColor("#143C36"))
        subtitle_style = ParagraphStyle("atama_subtitle", parent=styles["Heading2"], fontName=font_bold, fontSize=11, leading=14, alignment=TA_CENTER, spaceAfter=14)
        body_style = ParagraphStyle("atama_body", parent=styles["BodyText"], fontName=font, fontSize=10, leading=14, alignment=TA_LEFT, spaceAfter=10)
        foot_style = ParagraphStyle("atama_foot", parent=styles["BodyText"], fontName=font, fontSize=8, leading=11, textColor=colors.HexColor("#56615C"))

        doc_pdf = SimpleDocTemplate(str(tmp_pdf), pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=18*mm, bottomMargin=18*mm)
        story = [
            Paragraph(mapping["{{SIRKET_ADI}}"], title_style),
            Paragraph("ATAMA / GÖREV DEĞİŞİKLİĞİ BİLDİRİMİ", subtitle_style),
            Paragraph(f"Sevgili {isim},", body_style),
            Paragraph(
                f"Ekibimizdeki başarılı kariyer yolculuğunuzda; hedeflerinizi ve kişisel gelişiminizi ön planda "
                f"tutan çalışmalarınız sonucunda {mapping['{{MAGAZA}}']} Mağazasında {mapping['{{TARIH}}']} "
                f"tarihinde \u201c{mapping['{{YENI_POZISYON}}']}\u201d pozisyonuna atamanız gerçekleştirilmiştir.",
                body_style,
            ),
            Paragraph(
                "Göstermiş olduğunuz emek ile hem kendi, hem de ekibimizin hedeflerinin gerçekleşmesinde "
                "sağladığınız değerli katkılar için sizi kutluyor, yeni görevinizde başarılar diliyoruz.",
                body_style,
            ),
            Spacer(1, 6*mm),
            Paragraph("Birlikte başarılarımızın katlanarak artması dileğiyle...", body_style),
            Spacer(1, 10*mm),
            Paragraph("İnsan Kaynakları Direktörlüğü", ParagraphStyle("imza", parent=body_style, alignment=TA_CENTER)),
            Spacer(1, 14*mm),
            Paragraph(f"Önceki Pozisyon: {mapping['{{ONCEKI_POZISYON}}']} — {mapping['{{ONCEKI_MAGAZA}}']}", foot_style),
            Paragraph(f"Atama Tarihi: {mapping['{{TARIH}}']}    |    Onaylayan: {mapping['{{ONAYLAYAN}}']}", foot_style),
        ]
        doc_pdf.build(story)
        if not tmp_pdf.is_file() or tmp_pdf.stat().st_size < 1000 or not tmp_pdf.read_bytes().startswith(b"%PDF-"):
            raise IOError("Atama PDF dosyası oluşturulamadı.")
        make_outlook_safe_pdf(tmp_pdf)

        os.replace(tmp_docx, final_docx)
        os.replace(tmp_pdf, final_pdf)

    return {"docx": str(final_docx), "pdf": str(final_pdf)}
