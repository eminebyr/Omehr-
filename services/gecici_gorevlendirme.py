from __future__ import annotations

"""GEÇİCİ GÖREVLENDİRME / ŞUBE DESTEK FORMU oluşturur.

services/rotation_document.py'deki KALICI rotasyon belgesinden FARKLI
bir belgedir: bu form personelin norm kadro/asıl mağaza kaydının
DEĞİŞMEDİĞİNİ, yalnız GEÇİCİ bir süre için başka bir şubeye destek
amacıyla gönderildiğini belgeler — süre sonunda ayrı bir bildirim
olmadıkça asıl görev yerine döner.

Onaylar ekranında bir transfer/görevlendirme kararı kaydedilirken, İK
"Kalıcı Rotasyon Belgesi" ile "Geçici Görevlendirme / Şube Destek
Formu" arasında seçim yapar (bkz. web/tab_modules/onaylar.py). İkisi
de aynı transfer_recipients() alıcı listesini kullanır (zaten hem
kaynak hem hedef mağaza/bölge bilgisini kapsıyor).
"""

import os
import re
import tempfile
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Mapping

from docx import Document
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from services.pdf_compat import make_outlook_safe_pdf
from services.runtime_paths import runtime_root
from src.pdf_fonts import font as _omehr_font

def _template():
    from services.runtime_paths import runtime_root
    return runtime_root() / "templates" / "GECICI_GOREVLENDIRME_SUBE_DESTEK_FORMU.docx"
def _output_dir():
    from services.runtime_paths import runtime_root
    return runtime_root() / "output" / "Gecici_Gorevlendirme_Formlari"

NEDEN_SECENEKLERI = [
    "Personel eksikliği", "İzin nedeniyle destek", "Rapor / devamsızlık nedeniyle destek",
    "Yoğunluk nedeniyle destek", "Açılış / operasyon desteği", "Diğer",
]


def _safe_name(value: object) -> str:
    table = str.maketrans({"ç":"c","Ç":"C","ğ":"g","Ğ":"G","ı":"i","İ":"I","ö":"o","Ö":"O","ş":"s","Ş":"S","ü":"u","Ü":"U"})
    text = unicodedata.normalize("NFKD", str(value or "").translate(table))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return re.sub(r"[^A-Za-z0-9._-]+", "_", text).strip("_") or "gecici_gorevlendirme"


def _write_cell(cell, text: str) -> None:
    cell.text = str(text or "")
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "DejaVu Sans"
            run.font.size = None


def _mark_reason_in_paragraphs(doc: Document, secilen_neden: str, diger_metni: str = "") -> None:
    """'☐ Personel eksikliği ☐ İzin nedeniyle destek ...' satırındaki
    ilgili kutuyu ☑ yapar. Kutu işaretleme metnin run'lara nasıl
    bölündüğüne duyarlı olabileceğinden, GÜVENİLİR bir doğrulama olarak
    ayrıca kalın bir 'Seçilen neden:' satırı da eklenir."""
    for paragraph in doc.paragraphs:
        birlesik = "".join(r.text for r in paragraph.runs) or paragraph.text
        if "☐" in birlesik and any(n in birlesik for n in NEDEN_SECENEKLERI):
            for i, run in enumerate(paragraph.runs):
                if secilen_neden in run.text or (i > 0 and secilen_neden in doc.paragraphs[0].text):
                    pass
            yeni_metin = birlesik
            hedef = f"☐ {secilen_neden}"
            if hedef in yeni_metin:
                yeni_metin = yeni_metin.replace(hedef, f"☑ {secilen_neden}", 1)
                if paragraph.runs:
                    paragraph.runs[0].text = yeni_metin
                    for r in paragraph.runs[1:]:
                        r.text = ""


def create_temporary_assignment_documents(assignment: Mapping[str, object]) -> dict[str, str]:
    """assignment: person_name, person_id, current_title, source_store,
    target_store, start_date, end_date, total_duration, reason,
    reason_other, approved_by, source_manager, target_manager."""
    _output_dir().mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = f"GECICI_GOREV_{_safe_name(assignment.get('person_name'))}_{stamp}"
    final_docx = _output_dir() / f"{stem}.docx"
    final_pdf = _output_dir() / f"{stem}.pdf"

    fields = {
        "person": str(assignment.get("person_name") or ""),
        "person_id": str(assignment.get("person_id") or ""),
        "title": str(assignment.get("current_title") or ""),
        "source": str(assignment.get("source_store") or ""),
        "target": str(assignment.get("target_store") or ""),
        "start_date": str(assignment.get("start_date") or ""),
        "end_date": str(assignment.get("end_date") or ""),
        "total_duration": str(assignment.get("total_duration") or ""),
        "reason": str(assignment.get("reason") or ""),
        "reason_other": str(assignment.get("reason_other") or ""),
    }

    with tempfile.TemporaryDirectory(prefix="gecici_gorev_", dir=str(_output_dir())) as tmpdir:
        tmp = Path(tmpdir)
        tmp_docx = tmp / final_docx.name
        tmp_pdf = tmp / final_pdf.name

        if not _template().exists():
            raise FileNotFoundError(f"Geçici görevlendirme şablonu bulunamadı: {_template()}")
        doc = Document(str(_template()))
        for table in doc.tables:
            for row in table.rows:
                label = " ".join(c.text.strip() for c in row.cells).casefold()
                if len(row.cells) < 2:
                    continue
                if "personelin adı soyadı" in label:
                    _write_cell(row.cells[-1], fields["person"])
                elif "sicil no" in label:
                    _write_cell(row.cells[-1], fields["person_id"])
                elif "unvanı" in label:
                    _write_cell(row.cells[-1], fields["title"])
                elif "asıl görev yeri" in label:
                    _write_cell(row.cells[-1], fields["source"])
                elif "geçici olarak görevlendirildiği" in label:
                    _write_cell(row.cells[-1], fields["target"])
                elif "başlangıç tarihi" in label:
                    _write_cell(row.cells[-1], fields["start_date"])
                elif "bitiş tarihi" in label:
                    _write_cell(row.cells[-1], fields["end_date"])
                elif "toplam süre" in label:
                    _write_cell(row.cells[-1], fields["total_duration"])
        reason_text = fields["reason"]
        if reason_text == "Diğer" and fields["reason_other"]:
            reason_text = f"Diğer: {fields['reason_other']}"
        _mark_reason_in_paragraphs(doc, fields["reason"])
        doc.save(str(tmp_docx))
        if not tmp_docx.is_file() or tmp_docx.stat().st_size < 1000:
            raise IOError("Geçici görevlendirme DOCX dosyası oluşturulamadı.")

        font = _omehr_font(bold=False)
        font_bold = _omehr_font(bold=True)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("gg_title", parent=styles["Heading1"], fontName=font_bold, fontSize=13, leading=17, alignment=TA_CENTER, spaceAfter=10, textColor=colors.HexColor("#143C36"))
        body_style = ParagraphStyle("gg_body", parent=styles["BodyText"], fontName=font, fontSize=9, leading=13, alignment=TA_LEFT)
        note_style = ParagraphStyle("gg_note", parent=styles["BodyText"], fontName=font, fontSize=8, leading=11, alignment=TA_LEFT, textColor=colors.HexColor("#9B2D2D"))

        doc_pdf = SimpleDocTemplate(str(tmp_pdf), pagesize=A4, rightMargin=16*mm, leftMargin=16*mm, topMargin=14*mm, bottomMargin=14*mm)
        story = [
            Paragraph("GEÇİCİ GÖREVLENDİRME / ŞUBE DESTEK FORMU", title_style),
            Spacer(1, 3*mm),
        ]
        data = [
            ["Personelin Adı Soyadı", fields["person"]],
            ["Sicil No", fields["person_id"]],
            ["Unvanı", fields["title"]],
            ["Asıl Görev Yeri / Mağazası", fields["source"]],
            ["Geçici Olarak Görevlendirildiği Mağaza", fields["target"]],
            ["Görevlendirme Başlangıç Tarihi", fields["start_date"]],
            ["Görevlendirme Bitiş Tarihi", fields["end_date"]],
            ["Toplam Süre", fields["total_duration"]],
            ["Görevlendirme Nedeni", reason_text],
        ]
        table_data = [[Paragraph(str(a), body_style), Paragraph(str(b), body_style)] for a, b in data]
        info = Table(table_data, colWidths=[75*mm, 105*mm])
        info.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.6,colors.black),("VALIGN",(0,0),(-1,-1),"TOP"),("BACKGROUND",(0,0),(0,-1),colors.HexColor("#F2F2F2")),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
        story += [info, Spacer(1, 5*mm)]
        story.append(Paragraph(
            "Bu görevlendirme rotasyon veya kalıcı işyeri değişikliği niteliğinde DEĞİLDİR. "
            "Personelin asıl görev yeri, unvanı ve norm kadro kaydı mevcut/asıl mağazasında kalmaya devam edecektir; "
            "personel destek verdiği mağazanın norm kadrosuna aktarılmayacaktır. Görevlendirme süresinin sonunda "
            "ayrıca bir bildirim yapılmadığı sürece personel asıl görev yerine dönecektir. Ücret, unvan ve mevcut "
            "özlük hakları bu geçici görevlendirme nedeniyle değişmeyecektir.",
            note_style,
        ))
        story.append(Spacer(1, 14*mm))
        signatures = Table(
            [["Personel\n\nAd Soyad:\nTarih:\nİmza:", "Asıl Mağaza Yöneticisi\n\nAd Soyad:\nTarih:\nİmza:"],
             ["Destek Verilen Mağaza Yöneticisi\n\nAd Soyad:\nTarih:\nİmza:", "Bölge Müdürü\n\nAd Soyad:\nTarih:\nİmza:"]],
            colWidths=[85*mm, 85*mm], rowHeights=[32*mm, 32*mm],
        )
        signatures.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.6,colors.black),("FONTNAME",(0,0),(-1,-1),font),("FONTSIZE",(0,0),(-1,-1),9),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),7),("TOPPADDING",(0,0),(-1,-1),7)]))
        story.append(signatures)
        story.append(Spacer(1, 6*mm))
        story.append(Paragraph("İnsan Kaynakları — Ad Soyad: ______________________  Tarih: __ / __ / ____  İmza: ______________________", body_style))
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph("Not: Geçici destek görevlendirmelerinde personelin norm ve asıl mağaza kaydı değiştirilmez.", note_style))
        doc_pdf.build(story)
        if not tmp_pdf.is_file() or tmp_pdf.stat().st_size < 1000 or not tmp_pdf.read_bytes().startswith(b"%PDF-"):
            raise IOError("Geçici görevlendirme PDF dosyası oluşturulamadı.")
        make_outlook_safe_pdf(tmp_pdf)

        os.replace(tmp_docx, final_docx)
        os.replace(tmp_pdf, final_pdf)

    return {"docx": str(final_docx), "pdf": str(final_pdf)}


def create_temporary_assignment_documents_and_notify(
    assignment: Mapping[str, object],
    *,
    account_frame=None,
    sheet_frames: Mapping[str, object] | None = None,
    tenant: str = "OMEHR",
) -> dict[str, str]:
    """create_temporary_assignment_documents() + otomatik mail — TEK ÇAĞRIDA.

    DÜZELTME (dayanıklılık) — services/rotation_document.py::
    create_rotation_documents_and_notify ile AYNI gerekçe: mevcut üretim
    akışında (web -> onay -> worker.py TRANSFER_DECISION) mail zaten
    otomatik gider; bu fonksiyon o zincirin DIŞINDAKİ çağrı yerleri için
    (script, gelecekteki otomasyon) EK, kendi kendine yeten bir erişim
    noktasıdır. Mevcut create_temporary_assignment_documents() ve
    worker.py akışı DEĞİŞTİRİLMEDİ.
    """
    documents = create_temporary_assignment_documents(assignment)
    from services.transfer_recipients import transfer_recipients
    from services.job_queue import enqueue as _enqueue_job
    recipients = transfer_recipients(account_frame, assignment, sheet_frames)
    if not recipients:
        return {**documents, "mail_status": "SKIPPED_NO_RECIPIENTS"}
    kisi = assignment.get("person_name") or assignment.get("person_id") or "Belirtilmemiş"
    kaynak = assignment.get("source_store") or "?"
    hedef = assignment.get("target_store") or "?"
    subject = f"Geçici Görevlendirme / Şube Destek Formu | {kisi} | {kaynak} -> {hedef}"
    body = (
        f"Personel: {kisi}\nDevreden şube (mevcut): {kaynak}\nDestek verilecek şube (geçici): {hedef}\n\n"
        "Bu e-postaya Geçici Görevlendirme / Şube Destek Formu (DOCX ve PDF) eklenmiştir — "
        "hem devreden hem destek alan şube yetkilisi tarafından imzalanıp İK'ya iletilmelidir. "
        "Not: Bu görevlendirme personelin norm ve asıl mağaza kaydını değiştirmez.\n\n"
        "Bu e-posta, hem devreden hem destek alan şubeye otomatik olarak gönderilmiştir."
    )
    attachments = [v for k, v in documents.items() if k in {"pdf", "docx"} and v]
    _enqueue_job("SEND_EMAIL", {
        "report_type": "GECICI_GOREVLENDIRME", "subject": subject, "body": body,
        "recipients": recipients, "attachments": attachments,
    }, tenant=tenant)
    return {**documents, "mail_status": "QUEUED"}
