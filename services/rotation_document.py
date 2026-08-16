from __future__ import annotations

import os
import re
import tempfile
import unicodedata
from datetime import datetime
from pathlib import Path
from services.runtime_paths import runtime_root
from typing import Mapping

from docx import Document
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from services.pdf_compat import make_outlook_safe_pdf

def _template():
    from services.runtime_paths import runtime_root
    return runtime_root() / "templates" / "ROTASYON_BELGESI_SABLONU.docx"
def _output_dir():
    from services.runtime_paths import runtime_root
    return runtime_root() / "output" / "Rotasyon_Belgeleri"


def _safe_name(value: object) -> str:
    table = str.maketrans({"ç":"c","Ç":"C","ğ":"g","Ğ":"G","ı":"i","İ":"I","ö":"o","Ö":"O","ş":"s","Ş":"S","ü":"u","Ü":"U"})
    text = unicodedata.normalize("NFKD", str(value or "").translate(table))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return re.sub(r"[^A-Za-z0-9._-]+", "_", text).strip("_") or "rotasyon"


def _atomic_replace(temp_path: Path, final_path: Path) -> None:
    final_path.parent.mkdir(parents=True, exist_ok=True)
    os.replace(temp_path, final_path)


from src.pdf_fonts import font as _basdas_font


def _font_name() -> str:
    """Rotasyon belgesi PDF fontu — engine_core'un TEK Türkçe-doğrulamalı font
    kayıt noktasını (src/pdf_fonts.font) kullanır. Önceden burada AYRI ve
    daha zayıf bir kayıt vardı (Türkçe glif doğrulaması yok, sadece regular
    ağırlık, düşük boyut eşiği) — kaldırıldı, tekilleştirildi (bkz.
    FONT_TURKCE_DOGRULAMA.md)."""
    return _basdas_font(bold=False)


def _font_name_bold() -> str:
    return _basdas_font(bold=True)


def _write_cell(cell, text: str) -> None:
    cell.text = str(text or "")
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "DejaVu Sans"
            run.font.size = None


def create_rotation_documents(transfer: Mapping[str, object]) -> dict[str, str]:
    """Create a filled DOCX and PDF using temporary files, then atomically publish them."""
    _output_dir().mkdir(parents=True, exist_ok=True)
    tid = transfer.get("id", "")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = f"ROTASYON_{tid}_{_safe_name(transfer.get('person_name'))}_{stamp}"
    final_docx = _output_dir() / f"{stem}.docx"
    final_pdf = _output_dir() / f"{stem}.pdf"

    fields = {
        "person": str(transfer.get("person_name") or ""),
        "person_id": str(transfer.get("person_id") or ""),
        "source": str(transfer.get("source_store") or ""),
        "target": str(transfer.get("target_store") or ""),
        "current_title": str(transfer.get("current_title") or ""),
        "target_title": str(transfer.get("target_title") or ""),
        "reason": str(transfer.get("reason") or ""),
        "planned_date": str(transfer.get("planned_date") or ""),
        "source_region": str(transfer.get("region") or ""),
        "target_region": str(transfer.get("target_region") or ""),
        "decision_note": str(transfer.get("decision_note") or ""),
    }

    with tempfile.TemporaryDirectory(prefix="rotation_", dir=str(_output_dir())) as tmpdir:
        tmp = Path(tmpdir)
        tmp_docx = tmp / final_docx.name
        tmp_pdf = tmp / final_pdf.name

        if not _template().exists():
            raise FileNotFoundError(f"Rotasyon şablonu bulunamadı: {_template()}")
        doc = Document(str(_template()))
        for table in doc.tables:
            for row in table.rows:
                label = " ".join(c.text.strip() for c in row.cells).casefold()
                if len(row.cells) < 2:
                    continue
                if "adı – soyadı" in label or "adı - soyadı" in label:
                    _write_cell(row.cells[-1], fields['person'])
                elif "asıl görevlendirme yeri" in label:
                    _write_cell(row.cells[-1], f"{fields['source']} / {fields['current_title']}")
                elif "değişikliği nedeni" in label:
                    _write_cell(row.cells[-1], fields["reason"])
                elif "başlangıç tarihi" in label:
                    _write_cell(row.cells[-1], fields["planned_date"])
                elif "yeni görev yeri" in label:
                    _write_cell(row.cells[-1], f"{fields['target']} / {fields['target_title']}")
        doc.save(str(tmp_docx))
        if not tmp_docx.is_file() or tmp_docx.stat().st_size < 1000:
            raise IOError("Rotasyon DOCX dosyası oluşturulamadı.")

        font = _font_name()
        font_bold = _font_name_bold()
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("rot_title", parent=styles["Heading1"], fontName=font_bold, fontSize=12, leading=15, alignment=TA_CENTER, spaceAfter=6)
        body_style = ParagraphStyle("rot_body", parent=styles["BodyText"], fontName=font, fontSize=9, leading=12, alignment=TA_LEFT)
        doc_pdf = SimpleDocTemplate(str(tmp_pdf), pagesize=A4, rightMargin=14*mm, leftMargin=14*mm, topMargin=12*mm, bottomMargin=12*mm)
        story = [Paragraph("OMEHR İş Gücü Yönetimi ve Karar Destek Platformu", title_style)]
        header = Table([["Doküman Adı: ROTASYON BELGESİ", "İK.FR.004"]], colWidths=[145*mm, 35*mm])
        header.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.7,colors.black),("FONTNAME",(0,0),(-1,-1),font),("FONTSIZE",(0,0),(-1,-1),9),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#F2F2F2"))]))
        story += [header, Spacer(1,5*mm)]
        data = [
            ["Görev yeri değişikliği yapılan çalışanın Adı – Soyadı", fields['person']],
            ["Asıl Görevlendirme Yeri/Görevi", f"{fields['source']} / {fields['current_title']}"],
            ["Devreden Bölge Sorumlusu", fields["source_region"]],
            ["Görev yeri Değişikliği Nedeni", fields["reason"]],
            ["Görev yeri Değişikliği Başlangıç Tarihi", fields["planned_date"]],
            ["Yeni Görev Yeri/Görevi", f"{fields['target']} / {fields['target_title']}"],
            ["Devralan Bölge Sorumlusu", fields["target_region"]],
        ]
        table_data = [[Paragraph(str(a), body_style), Paragraph(str(b), body_style)] for a,b in data]
        info = Table(table_data, colWidths=[75*mm,105*mm])
        info.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.6,colors.black),("VALIGN",(0,0),(-1,-1),"TOP"),("BACKGROUND",(0,0),(0,-1),colors.HexColor("#F2F2F2")),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
        story += [info, Spacer(1,6*mm)]
        story.append(Paragraph("Görevlendirilen Namına: Yukarıda belirtilen tarihten itibaren görevlendirme talebini kabul ediyorum. Görevlendirildiğim çalışma alanı içindeki tüm İş Sağlığı ve Güvenliği kurallarına uyacağımı beyan ederim.", body_style))
        story.append(Spacer(1,4*mm))
        story.append(Paragraph("Görevlendiren Namına: Yukarıda belirtilen tarihten itibaren görevlendirilen personelin görev süresinde iş sağlığı ve güvenliği kurallarına uygun şekilde çalıştırılacağını beyan ederim.", body_style))
        story.append(Spacer(1,18*mm))
        signatures = Table([["Görevlendirme yapılan çalışan\n\nAd-Soyad:\nTarih:\nİmza:", "Görevlendirme tebliğ eden\n\nAd-Soyad:\nTarih:\nİmza:"]], colWidths=[90*mm,90*mm], rowHeights=[55*mm])
        signatures.setStyle(TableStyle([("GRID",(0,0),(-1,-1),0.6,colors.black),("FONTNAME",(0,0),(-1,-1),font),("FONTSIZE",(0,0),(-1,-1),9),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),7),("TOPPADDING",(0,0),(-1,-1),7)]))
        story.append(signatures)
        doc_pdf.build(story)
        if not tmp_pdf.is_file() or tmp_pdf.stat().st_size < 1000 or not tmp_pdf.read_bytes().startswith(b"%PDF-"):
            raise IOError("Rotasyon PDF dosyası oluşturulamadı.")
        # DÜZELTME: bu PDF, yönetici/admin raporlarıyla AYNI ReportLab TTF
        # font kaydını (src/pdf_fonts.font) kullanıyor — dolayısıyla AYNI
        # Windows Outlook/Edge font alt-küme render riskini taşıyor. Üstelik
        # bu, tam olarak Outlook ile devreden/devralan şubeye gönderilen
        # belge türü — sorunun ilk fark edildiği senaryonun kendisi. Diğer
        # raporlarla tutarlı olması için aynı image-only dönüşüm uygulanır.
        make_outlook_safe_pdf(tmp_pdf)

        _atomic_replace(tmp_docx, final_docx)
        _atomic_replace(tmp_pdf, final_pdf)

    return {"docx": str(final_docx), "pdf": str(final_pdf)}
